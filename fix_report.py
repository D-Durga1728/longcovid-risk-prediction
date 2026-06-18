import docx
import shutil
from datetime import datetime

src = r'C:\Users\durga\OneDrive\Desktop\Practicum\DUR#12839\Report.docx'
bak = src.replace('.docx', f'_backup_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
shutil.copy2(src, bak)
print(f"Backup: {bak}")

doc = docx.Document(src)
changes = []


def replace_in_para(para, old, new):
    combined = ''.join(r.text for r in para.runs)
    if old not in combined:
        return False
    replaced = combined.replace(old, new)
    if para.runs:
        para.runs[0].text = replaced
        for r in para.runs[1:]:
            r.text = ''
    return True


# ── 1. LR para (53): correct AUC, remove SMOTE, add DeLong + calibration ──
p53 = doc.paragraphs[53]
if p53.text.strip().startswith("The comparative analysis of the baseline and advanced Logistic Regression"):
    new53 = (
        "The final deployed model is a calibrated, tuned Logistic Regression (C = 0.01), achieving a "
        "test AUC of 0.888 (95% CI 0.884-0.892, bootstrap). A leakage-free pipeline was enforced: the "
        "imputer and scaler were fit on the training partition only and applied to the held-out test set. "
        "Class imbalance (12.3% mortality) was addressed via cost-sensitive learning (class_weight='balanced'). "
        "SMOTE was evaluated and rejected because it produced an identical AUC (0.888) while fabricating "
        "132,810 synthetic patients, introducing unnecessary distributional distortion. Calibration was poor "
        "before correction (Brier 0.135, ECE 0.195) and substantially improved after isotonic calibration "
        "(Brier 0.078, count-weighted ECE 0.002), meaning the output probability can be interpreted directly "
        "as a mortality risk estimate. Decision Curve Analysis confirms positive net benefit over treat-all "
        "and treat-none strategies. DeLong testing showed no model significantly outperformed Logistic "
        "Regression (p > 0.05 for GB, XGBoost, Stacking), supporting deployment of the interpretable model. "
        "Odds Ratios (statsmodels, per +1 SD, 95% CI) identify pneumonia (OR 2.65), age (OR 2.33), diabetes "
        "(OR 1.15) and obesity (OR 1.10) as the largest risk drivers; cardiovascular and asthma are not "
        "statistically significant."
    )
    if p53.runs:
        p53.runs[0].text = new53
        for r in p53.runs[1:]:
            r.text = ''
    changes.append("Para 53: rewrote LR section — AUC 0.888, removed SMOTE, added DeLong + calibration numbers")

# ── 2. RF para (57): fix AUC values ──
p57 = doc.paragraphs[57]
for old, new in [
    ("AUC = 89.64%", "AUC = 86.6%"),
    ("AUC = 93.80%", "AUC = 86.7%"),
    ("AUC = 93.80", "AUC = 86.7"),
    ("minority recall to 0.93", "minority-class recall improved"),
    ("Brier Score of 0.098", "Brier Score of 0.098 (uncalibrated; calibration recommended)"),
]:
    if replace_in_para(p57, old, new):
        changes.append(f"Para 57 RF: '{old}' -> '{new}'")

# ── 3. GB para (61): fix AUC values ──
p61 = doc.paragraphs[61]
for old, new in [
    ("AUC = 94.02%", "AUC = 88.8%"),
    ("AUC - 93.85%", "AUC = 88.8%"),
    ("Brier score (0.097)", "Brier score (0.097, uncalibrated)"),
]:
    if replace_in_para(p61, old, new):
        changes.append(f"Para 61 GB: '{old}' -> '{new}'")

# ── 4. SHAP para (65): fix dominant feature ──
p65 = doc.paragraphs[65]
shap_fixes = [
    (
        "Patient type is the dominant feature showing that severity of disease at presentation is "
        "the main determinant of adverse outcomes, aligning with triage dynamics observed in real-world. "
        "Age is the second-best contributor, showing its well-established link with immunological problems "
        "after covid and progression toward severe complications in respiration. Exposure related features "
        "and pneumonia further show substantial influence in prediction, showing the role of acute "
        "respiratory inclusion and intensity of transmission in shaping risk of mortality.",
        "Pneumonia is the dominant feature (Tree-SHAP global importance 44.2%), confirming that acute "
        "respiratory involvement is the primary determinant of mortality risk. Age is the second-largest "
        "contributor (33.0%), consistent with its well-established link with immunological decline and "
        "severe post-COVID complications. Sex (female, protective; 8.5%), diabetes (5.6%) and hypertension "
        "(4.5%) follow in importance. Cardiovascular disease and asthma show minimal contribution (0.3% "
        "each), consistent with their non-significant odds ratios. The 9-feature model deliberately "
        "excludes intubation and ICU status — these are post-admission treatment outcomes, not "
        "admission-time predictors, and their inclusion would constitute target leakage."
    ),
]
for old, new in shap_fixes:
    if replace_in_para(p65, old, new):
        changes.append("Para 65: replaced 'Patient type' with correct pneumonia-dominant Tree-SHAP narrative + leakage note")

# ── 5. Reference fix: Esquivel -> Guzman-Esquivel ──
p92 = doc.paragraphs[92]
for old, new in [
    ("Esquivel, J.G. et al.", "Guzman-Esquivel, J.G. et al."),
    ("Esquivel, J.G.", "Guzman-Esquivel, J.G."),
]:
    if replace_in_para(p92, old, new):
        changes.append(f"Para 92 ref: '{old}' -> '{new}'")
        break

# ── 6. Add Cordelli 2024 to reference list if missing ──
cordelli_exists = any("Cordelli" in p.text for p in doc.paragraphs)
if not cordelli_exists:
    cordelli_ref = (
        "Cordelli, E. et al. (2024) 'Predicting pulmonary sequelae after COVID-19 using machine learning: "
        "model performance and clinical interpretability'. [Reference to be verified against original publication]."
    )
    p_dozier = doc.paragraphs[91]
    new_para = doc.add_paragraph(cordelli_ref, style=p_dozier.style)
    p_dozier._element.addprevious(new_para._element)
    # Remove the duplicate appended at end
    last = doc.paragraphs[-1]
    if "Cordelli" in last.text:
        last._element.getparent().remove(last._element)
    changes.append("Added Cordelli 2024 to reference list before Dozier")

# ── 7. Cost-benefit: label illustrative ──
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if "$2.93 B" in t:
        if replace_in_para(p, "$2.93 B", "$2.93B (illustrative scenario, not empirically validated)"):
            changes.append(f"Para {i}: labelled $2.93B as illustrative")
    if "$5.87 B" in t:
        if replace_in_para(p, "$5.87 B", "$5.87B (illustrative scenario, not empirically validated)"):
            changes.append(f"Para {i}: labelled $5.87B as illustrative")
    if "$8.80 B" in t:
        if replace_in_para(p, "$8.80 B", "$8.80B (illustrative scenario, not empirically validated)"):
            changes.append(f"Para {i}: labelled $8.80B as illustrative")

# ── 8. Supervisor fix ──
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if "Martin Crane" in t and "Tai Tan Mai" not in t:
        if replace_in_para(p, "Dr Martin Crane", "Dr Martin Crane & Dr Tai Tan Mai"):
            changes.append(f"Para {i}: added Dr Tai Tan Mai as co-supervisor")
        elif replace_in_para(p, "Martin Crane", "Martin Crane & Dr Tai Tan Mai"):
            changes.append(f"Para {i}: added Dr Tai Tan Mai as co-supervisor (alt)")

# ── 9. Abstract: remove overstated AUC if present ──
p3 = doc.paragraphs[3]
for old, new in [
    ("AUC = 0.93", "AUC = 0.888"),
    ("AUC=0.93", "AUC=0.888"),
    ("AUC 0.93", "AUC 0.888"),
]:
    if replace_in_para(p3, old, new):
        changes.append(f"Abstract: fixed AUC {old} -> {new}")

doc.save(src)
print(f"\nSaved to {src}")
print(f"\nChanges applied ({len(changes)}):")
for c in changes:
    print(f"  OK  {c}")
if not changes:
    print("  (No matching text found for any replacement — check paragraph indices)")

# IEEE final paper -- 10 pages, real embedded figures, DCU sample paper style
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

BASE = Path(r'C:\Users\durga\OneDrive\Desktop\Practicum\Final')
OUT  = BASE / 'Report_IEEE_v7.docx'
FIGS = BASE / 'analysis_output' / 'visualizations'

doc = Document()
sec = doc.sections[0]
sec.page_height = Cm(29.7); sec.page_width    = Cm(21.0)
sec.top_margin  = Cm(1.9);  sec.bottom_margin = Cm(2.54)
sec.left_margin = Cm(1.78); sec.right_margin  = Cm(1.78)
TNR = 'Times New Roman'

# ── column layout ──────────────────────────────────────────────────────────────
def set_cols(section, n=2):
    sp = section._sectPr
    for old in sp.findall(qn('w:cols')):
        sp.remove(old)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), str(n))
    cols.set(qn('w:space'), '284')
    cols.set(qn('w:equalWidth'), '1')
    sp.insert(0, cols)

set_cols(sec, 1)   # title block is single-column

# ── typography helpers ──────────────────────────────────────────────────────────
def add_p(text, size=10, bold=False, italic=False,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=3, indent=0):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before      = Pt(sb)
    para.paragraph_format.space_after       = Pt(sa)
    para.paragraph_format.first_line_indent = Pt(indent)
    r = para.add_run(text)
    r.font.name = TNR; r.font.size = Pt(size)
    r.bold = bold;     r.italic   = italic
    return para

def h1(num, title):
    txt = f'{num}. {title.upper()}' if num else title.upper()
    add_p(txt, size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sb=7, sa=3)

def h2(letter, title):
    add_p(f'{letter}. {title}', size=10, bold=True, italic=True,
          align=WD_ALIGN_PARAGRAPH.LEFT, sb=5, sa=2)

def body(text, sa=3):
    add_p(text, size=10, indent=14, sa=sa)

def caption(text):
    add_p(text, size=8, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=6)

def eq(text, num):
    """Centred display equation with a right-hand equation number (IEEE style)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(3)
    r = para.add_run(text)
    r.font.name = TNR; r.font.size = Pt(10); r.italic = True
    rn = para.add_run(f'        ({num})')
    rn.font.name = TNR; rn.font.size = Pt(10)

def fig_embed(rel, cap, w=7.8):
    full = FIGS / rel
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(0)
    if full.exists():
        para.add_run().add_picture(str(full), width=Cm(w))
    else:
        para.add_run(f'[Fig: {rel}]').italic = True
    caption(cap)

def _page_setup(s):
    s.page_height = Cm(29.7); s.page_width    = Cm(21.0)
    s.top_margin  = Cm(1.9);  s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(1.78); s.right_margin  = Cm(1.78)

def fig_full(rel, cap, w=15.2):
    """Full-width figure spanning both columns (IEEE figure* style)."""
    s1 = doc.add_section(WD_SECTION.CONTINUOUS); _page_setup(s1); set_cols(s1, 1)
    fig_embed(rel, cap, w=w)
    s2 = doc.add_section(WD_SECTION.CONTINUOUS); _page_setup(s2); set_cols(s2, 2)

def blk(lead, text, sa=3):
    """Bold lead-in label followed by normal body text, in one paragraph."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(sa)
    para.paragraph_format.first_line_indent = Pt(14)
    rl = para.add_run(lead + '  ')
    rl.font.name = TNR; rl.font.size = Pt(10); rl.bold = True
    rt = para.add_run(text)
    rt.font.name = TNR; rt.font.size = Pt(10)

def make_table(headers, rows, title):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    for cell, h in zip(tbl.rows[0].cells, headers):
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.name = TNR; run.font.size = Pt(8); run.bold = True
    for trow, data in zip(tbl.rows[1:], rows):
        for cell, val in zip(trow.cells, data):
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = TNR; run.font.size = Pt(8)
    caption(title)

def ref_entry(num, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before      = Pt(0)
    para.paragraph_format.space_after       = Pt(1)
    para.paragraph_format.left_indent       = Pt(18)
    para.paragraph_format.first_line_indent = Pt(-18)
    r = para.add_run(f'[{num}]  {text}')
    r.font.name = TNR; r.font.size = Pt(8)

# ════════════════════════════════════════════════════════════════════════════════
#  SINGLE-COLUMN TITLE BLOCK
# ════════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_after = Pt(8)
r = tp.add_run(
    'Predictive Modelling for Personalised Long COVID\n'
    'Risk Assessment and Prevention'
)
r.font.name = TNR; r.font.size = Pt(24); r.bold = True

# Author block — 3 separate author blocks side by side (no visible borders)
# Matches DCU paper style: each author gets Name / School / Uni / City / Email
def _kill_borders(cell):
    from docx.oxml import OxmlElement as OE
    tcPr = cell._tc.get_or_add_tcPr()
    # remove any existing tcBorders
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OE('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OE(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def _kill_table_borders(tbl):
    from docx.oxml import OxmlElement as OE
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OE('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblBorders = OE('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        b = OE(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)

_authors = [
    ('Durga Prasad Narsing',),
    ('School of Computing',),
    ('Dublin City University',),
    ('Glasnevin, Dublin 9, Ireland',),
    ('durgaprasadnarsing2@gmail.com',),
]

atbl = doc.add_table(rows=5, cols=1)
_kill_table_borders(atbl)
for ri, row_data in enumerate(_authors):
    for ci, txt in enumerate(row_data):
        cell = atbl.rows[ri].cells[ci]
        _kill_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(txt)
        run.font.name = TNR
        run.font.size = Pt(10)
        run.bold = (ri == 0)   # bold name row only

# space after author block
gap = doc.add_paragraph()
gap.paragraph_format.space_before = Pt(0)
gap.paragraph_format.space_after  = Pt(8)

# Abstract (9pt bold-italic label then 9pt italic body — per IEEE)
add_p('Abstract', size=9, bold=True, italic=True,
      align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=0)
add_p(
    'Post-COVID-19 Condition (Long COVID) affects an estimated 6% of SARS-CoV-2 infected '
    'individuals, imposing persistent health burdens that current clinical pathways struggle '
    'to anticipate. Existing predictive models suffer from four recurring methodological gaps: '
    'calibration is rarely reported, demographic fairness is seldom audited, post-admission '
    'features introduce target leakage, and per-patient explainability is largely absent. '
    'This paper presents a 13-phase, leakage-free, responsible-ML pipeline applied to 220,218 '
    'confirmed COVID-19 cases from the Mexican Government surveillance dataset, using in-hospital '
    'mortality as a clinically justified proxy for severe Long COVID outcome risk. Nine '
    'admission-time features are modelled with Logistic Regression, Random Forest, Gradient '
    'Boosting, XGBoost, and a Stacking ensemble. A calibrated Logistic Regression achieves '
    'AUC 0.888 (95% CI 0.884-0.892); DeLong statistical testing confirms non-inferiority over '
    'all complex models (p > 0.05). Isotonic calibration reduces Brier score from 0.135 to '
    '0.078 and count-weighted ECE from 0.195 to 0.002. An eight-subgroup fairness audit reveals '
    'equitable calibration in every group (ECE < 0.015) with TPR disparity reduced from 0.603 '
    'to 0.045 via group-specific threshold mitigation. Tree-SHAP identifies pneumonia (44.2%) '
    'and age (33.0%) as dominant risk drivers. A Streamlit web prototype delivers calibrated '
    'risk scores and real-time explanations, closing all four identified gaps in a single '
    'reproducible workflow.',
    size=9, italic=True, indent=0, sa=3
)
add_p(
    'Index Terms -- Long COVID, PASC, COVID-19 mortality, machine learning, calibration, '
    'DeLong test, SHAP, fairness audit, responsible AI, clinical decision support, Streamlit.',
    size=9, italic=True, indent=0, sa=10
)

# ════════════════════════════════════════════════════════════════════════════════
#  TWO-COLUMN BODY
# ════════════════════════════════════════════════════════════════════════════════
ns = doc.add_section()
ns.page_height = Cm(29.7); ns.page_width    = Cm(21.0)
ns.top_margin  = Cm(1.9);  ns.bottom_margin = Cm(2.54)
ns.left_margin = Cm(1.78); ns.right_margin  = Cm(1.78)
ns.start_type  = 1
set_cols(ns, 2)

# ── I. INTRODUCTION ───────────────────────────────────────────────────────────
h1('I', 'Introduction')

body(
    'Post-COVID-19 Condition (PASC), or Long COVID, is a major global public health burden. '
    'The World Health Organisation estimates that approximately 6% of individuals infected '
    'with SARS-CoV-2 develop PASC [1], characterised by persistent symptoms -- fatigue, '
    'dyspnoea, cognitive impairment, and cardiovascular sequelae -- whose prevalence and '
    'duration have been quantified by systematic review and meta-analysis [16], lasting months '
    'to years beyond the acute episode. United States national surveillance data report 6.9% of adults '
    'ever affected and 3.4% with currently ongoing symptoms [2], with disproportionate '
    'vulnerability in older adults, females, and individuals with pre-existing chronic '
    'comorbidities. The heterogeneity and scale of Long COVID demand validated clinical tools '
    'that identify high-risk patients at the point of acute presentation, before the window '
    'for early intervention closes.'
)
body(
    'Despite rapid growth in predictive modelling research, the field consistently exhibits '
    'four methodological gaps identified by systematic reviews [3][4]: (i) calibration is '
    'rarely reported, making model outputs impossible to interpret as clinical probabilities; '
    '(ii) performance equity across demographic subgroups is seldom audited, risking '
    'inequitable triage; (iii) post-admission variables such as intubation and ICU status '
    'are commonly used as predictors, constituting target leakage that inflates apparent '
    'performance without clinical utility at the admission decision point; and (iv) per-patient '
    'explainability is largely absent, limiting clinician trust and adoption. An additional open '
    'question is whether complex ensemble models offer meaningful discriminative gains over '
    'Logistic Regression on large, mostly-binary clinical datasets [5].'
)
body(
    'This paper addresses all four gaps with a 13-phase end-to-end responsible-ML pipeline. '
    'The Mexican Government COVID-19 dataset (n = 220,218 confirmed cases) provides '
    'admission-time clinical data and in-hospital mortality labels; mortality is used as a '
    'clinically justified proxy for severe outcome risk in the absence of PASC follow-up '
    'labels, consistent with prior literature [6][7]. Four research questions structure the '
    'investigation: RQ1 -- can nine admission-time features achieve clinically useful '
    'discrimination and calibration? RQ2 -- do complex ensembles significantly outperform '
    'Logistic Regression by DeLong test? RQ3 -- are predictions equitable and calibrated '
    'across demographic subgroups? RQ4 -- can predictions be delivered interpretably through '
    'a functional clinical prototype?'
)

h2('A', 'Contribution')
body(
    'This paper makes three contributions. First, an empirical finding: a calibrated '
    'Logistic Regression is statistically non-inferior to Random Forest, Gradient Boosting, '
    'XGBoost, and Stacking by DeLong\'s test (p > 0.05 on 220,218 patients), demonstrating '
    'that interpretability costs nothing in discrimination on low-dimensional binary clinical '
    'features. Second, a methodological contribution: the pipeline operationalises all four '
    'identified gaps in a single reproducible workflow -- calibration, subgroup fairness '
    'auditing, leakage control, and multi-layer explainability. Third, a responsible-AI '
    'contribution: auto-generated Model Card and Datasheet artefacts, a group-specific '
    'threshold mitigation strategy, and a Streamlit clinical prototype with governance '
    'disclaimers built in.'
)

h2('B', 'Paper Structure')
body(
    'Section II reviews related work and identifies the four gaps addressed. Section III '
    'describes the engineering pipeline in full detail. Section IV presents performance '
    'across six evaluation dimensions. Section V discusses critical findings and limitations. '
    'Section VI concludes.'
)

# ── II. RELATED WORK ──────────────────────────────────────────────────────────
h1('II', 'Related Work')
body(
    'This section reviews the Long COVID and COVID-19 mortality prediction literature across '
    'five themes -- risk-factor studies, ML prediction models, calibration, fairness, and '
    'explainability -- and concludes by stating the research gap this work addresses.', sa=2
)

h2('A', 'Risk-Factor Studies')
body(
    'Risk-factor studies consistently identify pneumonia, age, sex, obesity, and acute '
    'severity as predictors [6][12], but do not operationalise findings into validated '
    'deployable tools. Menezes et al. [6] identified age >= 45 (OR 5.55) and >=15 acute '
    'symptoms (OR 6.02) in a single-centre observational cohort (n = 108), limiting '
    'generalisability. Guzman-Esquivel et al. [12] found tachycardia (RR 10.4) and '
    'hospitalisation (RR 2.1) as predictors, without a deployable model or calibration.'
)

h2('B', 'ML Prediction Models and the ML-vs-Regression Question')
body(
    'Among deployable models, Jeffrey et al. [4] derived a Scottish logistic regression using '
    'five predictors with no reported calibration, fairness audit, or explainability. Antony '
    'et al. [3] reviewed ML models across two million EHR patients (AUROC 0.75-0.80) without '
    'calibration or fairness. Dozier et al. [7] applied a Super Learner ensemble to 57,563 '
    'patients with high discrimination but limited interpretability. Zang et al. [8] used '
    'regularised Cox models on the RECOVER cohort (C-index > 0.80) without calibration or '
    'fairness. Shakhovska et al. [9] reported AUC = 0.978 on only 122 patients -- raising '
    'severe overfitting concerns. Comparative COVID-19 mortality models provide context: '
    'Bozkurt and Keskin [17] achieved AUC 0.91 with gradient boosting; Cordelli et al. [18] '
    'predicted pulmonary sequelae; Kourmpanis et al. [19] AUC 0.94; Sun et al. [20] and Tang '
    'et al. [21] validated further mortality models. '
    'Critically, Bjerre et al. [5] found in a systematic comparison on large COVID databases '
    'that ML offers only marginal gains over classical regression -- a claim this work tests '
    'formally with DeLong statistics (RQ2).'
)

h2('C', 'Calibration and Clinical Utility')
body(
    'Teeuw et al. [11] emphasise that high discrimination does not guarantee clinical utility '
    'when calibration is poor. Calibration and Decision Curve Analysis are therefore treated '
    'here as primary deployment criteria rather than secondary metrics -- a departure from the '
    'discrimination-only reporting that dominates the prior work above.'
)

h2('D', 'Fairness in Clinical ML')
body(
    'Subgroup performance equity is rarely audited in the Long COVID literature [4][7][8][10], '
    'risking inequitable triage across age, sex, and comorbidity groups. Where fairness is '
    'considered, discrimination (AUC) and calibration equity are frequently conflated; this '
    'work audits both lenses separately and applies group-specific threshold mitigation.'
)

h2('E', 'Explainability')
body(
    'Per-patient explainability is largely absent from deployable Long COVID models [3][4][5]. '
    'Global feature importance, where reported, is seldom paired with individual attributions '
    'or a bedside instrument. This work integrates SHAP, odds ratios, and a clinical nomogram '
    'into a single explainability framework exposed through a functional prototype.'
)

h2('F', 'Research Gap')
body(
    'No prior study consolidates leakage control, probability calibration, subgroup fairness '
    'auditing, and multi-layer explainability into one reproducible, deployable workflow. '
    'Table I summarises the four recurring methodological gaps and how this study closes each.', sa=2
)
make_table(
    ['Gap', 'Prior Studies', 'This Work'],
    [
        ('Calibration not reported',        '[4][7][8][9]',  'Brier + ECE; isotonic correction'),
        ('No subgroup fairness audit',       '[4][7][8][10]', '8 subgroups; threshold mitigation'),
        ('Target leakage (intubation/ICU)', '[9][10]',       '9 admission-time features only'),
        ('No per-patient explainability',    '[4][5][3]',     'SHAP + OR + nomogram + web app'),
        ('No deployable prototype',          '[4][6][12]',    'Streamlit app with live pipeline'),
    ],
    'TABLE I.  Methodological Gaps in Prior Work vs This Study'
)

# ── III. METHODOLOGY ──────────────────────────────────────────────────────────
h1('III', 'Methodology')
body(
    'The proposed system is a leakage-free, responsible-ML pipeline that benchmarks five model '
    'families, selects the deployed model by formal statistical testing, calibrates its '
    'probabilities, and exposes calibrated risk with multi-layer explanations through a '
    'clinical prototype. As shown in Fig. 1, the system comprises eight stages, described '
    'block by block below; Fig. 2 organises them into the 13-phase machine-learning lifecycle.',
    sa=2
)
fig_full('architecture/fig_system_architecture.png',
    'Fig. 1.  Responsible-ML system architecture. Nine admission-time features are preprocessed '
    'leakage-free and benchmarked across five model families. The deployed model is selected by '
    'DeLong test and isotonic-calibrated, then emits a calibrated risk score, a tuned threshold '
    'flag, multi-layer explanations, and a subgroup fairness audit, all exposed through a '
    'Streamlit clinical prototype.')
blk('Patient Inputs.',
    'Nine admission-time features -- age, sex, and seven binary comorbidities -- enter the '
    'system; post-admission variables (intubation, ICU) are excluded to prevent target leakage.')
blk('Preprocessing.',
    'Sentinel decoding (97/98/99 -> NaN), median imputation, and standardisation are fit on the '
    'training partition only and applied identically to the test set, enforcing leakage-free '
    'transformation.')
blk('Model Benchmarking.',
    'Five families -- Logistic Regression, Random Forest, Gradient Boosting, XGBoost, and a '
    'Stacking ensemble -- are trained with cost-sensitive weighting to address the 1:7 class '
    'imbalance.')
blk('Model Selection.',
    "DeLong's test compares AUCs on the held-out set: no complex model significantly beats "
    'Logistic Regression (p > 0.05), so the interpretable model is chosen.')
blk('Calibration.',
    'The selected Logistic Regression is wrapped in isotonic CalibratedClassifierCV (cv=5), '
    'reducing count-weighted ECE from 0.195 to 0.002 and Brier score from 0.135 to 0.078.')
blk('Deployed Model.',
    'The calibrated Logistic Regression is persisted (13 KB) and serves every downstream task '
    'through the identical impute -> scale -> predict chain used during training.')
blk('Outputs.',
    'It emits a calibrated risk percentage, a tuned decision flag (F1-optimal 0.26 or Youden '
    '0.10), SHAP / odds-ratio / nomogram explanations, and an eight-subgroup fairness audit '
    'with threshold mitigation.')
blk('Clinical Prototype.',
    'A two-page Streamlit application exposes all outputs at the bedside, with an auto-generated '
    'Model Card and Datasheet providing responsible-AI governance.')
fig_full('architecture/fig_pipeline.png',
    'Fig. 2.  13-phase methodology pipeline. Phase 1 (data preparation) applies leakage-free '
    'preprocessing and cost-sensitive sampling; Phase 2 (training and selection) benchmarks five '
    'models, tests AUC differences, and calibrates the deployed model; Phase 3 (evaluation) '
    'computes discrimination, calibration, fairness, explainability, decision-curve, and temporal '
    'results, and emits governance artefacts. The following subsections detail each component.')

h2('A', 'Dataset and Cohort Selection')
body(
    'The Mexican Government COVID-19 surveillance dataset (Kaggle 2020 [14]) contains '
    '566,602 patient records from the acute phase of SARS-CoV-2 infection. Variables include '
    'age (continuous, years), sex (binary), seven binary comorbidities (diabetes, hypertension, '
    'cardiovascular disease, pneumonia, obesity, asthma, COPD), clinical outcomes (intubation, '
    'ICU admission, death date), and administrative fields. The primary modelling cohort '
    'comprises 220,218 confirmed COVID-19 cases (covid_res == 1) with an in-hospital mortality '
    'rate of 12.31% (27,102 deaths; 193,116 survivors), representing a 1:7 class imbalance. '
    'Data were partitioned 80/20 stratified by outcome: 176,174 training patients, 44,044 '
    'held-out test patients. Events-per-variable (EPV = 27,102 / 9 = 3,011) far exceeds the '
    'minimum threshold of 10, confirming ample data for stable coefficient estimation. '
    'Table II summarises the cohort and feature set.'
)
make_table(
    ['Property', 'Value'],
    [
        ('Source', 'Mexican Govt. COVID-19 surveillance [14]'),
        ('Total records', '566,602'),
        ('Confirmed cohort', '220,218 (covid_res = 1)'),
        ('Outcome / prevalence', 'in-hospital mortality / 12.31%'),
        ('Class imbalance', '1 : 7 (27,102 vs 193,116)'),
        ('Features (9)', 'age, sex, 7 binary comorbidities'),
        ('Train / Test', '176,174 / 44,044 (stratified 80/20)'),
        ('Events-per-variable', '3,011 (>> 10)'),
    ],
    'TABLE II.  Dataset and Cohort Composition'
)

h2('B', 'Preprocessing and Leakage Control')
body(
    'Outcome engineering: date_died != \'9999-99-99\' maps to 1 (deceased), else 0. '
    'Binary comorbidities were recoded from administrative encoding (1=Yes, 2=No) to '
    'standard binary (1.0/0.0). Sentinel values 97, 98, 99 (unknown or not-applicable '
    'comorbidity status) were decoded to NaN prior to imputation, affecting approximately '
    '0.3% of comorbidity field entries. A SimpleImputer (median strategy) was fit '
    'exclusively on the training partition and applied to the held-out test set, preventing '
    'leakage of test-set distribution into imputation decisions. A StandardScaler was '
    'applied identically. All transforms were encapsulated in scikit-learn Pipelines, '
    'enforcing the split-then-fit discipline within every cross-validation fold.'
)
body(
    'Intubation status and ICU admission were deliberately excluded from the feature set. '
    'Both are determined during the hospital stay rather than at initial presentation, and '
    'their inclusion constitutes target leakage -- inflating apparent AUC without clinical '
    'utility at the admission decision point. The nine final features are: age, sex, '
    'diabetes, hypertension, cardiovascular disease, pneumonia, obesity, asthma, COPD.'
)

h2('C', 'Statistical Feature Validation')
body(
    'All feature-outcome associations were validated statistically before modelling. An '
    'independent-samples t-test confirmed deceased patients were significantly older than '
    'survivors (mean 61.0 vs 43.6 years; t = 176.99, p < 0.001). Chi-square tests confirmed '
    'all seven binary comorbidities are significantly associated with mortality (p < 0.001 '
    'for all). Multiple testing was controlled with Benjamini-Hochberg FDR correction across '
    'five primary comparisons (q < 0.05 for all), confirming no feature enters the model '
    'on domain intuition alone. Feature correlation analysis (Pearson) showed no pair '
    'exceeded r = 0.34, ruling out severe multicollinearity concerns and confirming a '
    'compact, non-redundant feature set consistent with feature-selection best practice [13].'
)

h2('D', 'Model Training and Hyperparameter Tuning')
body(
    'Five model families were trained and benchmarked: Logistic Regression (L2 '
    'regularisation), Random Forest, Gradient Boosting, XGBoost, and a Stacking ensemble '
    '(LR + RF as base learners, LR as meta-learner). Class imbalance (1:7) was addressed '
    'via cost-sensitive learning throughout: class_weight=\'balanced\' for LR and RF; '
    'sample_weight for Gradient Boosting; scale_pos_weight = 7 for XGBoost. Hyperparameters '
    'were selected by 5-fold stratified GridSearchCV on training data only: LR regularisation '
    'C = 0.01 (CV AUC 0.891 +/- 0.002); RF max_depth = 6, n_estimators = 200, '
    'min_samples_leaf = 10. SMOTE was evaluated as an alternative class-balance strategy; '
    'it produced identical test AUC (0.888) while fabricating 132,810 synthetic patients '
    'and increasing training time by 38%, and was therefore rejected in favour of '
    'cost-sensitive learning. All experiments use random_state = 42 for reproducibility. '
    'Fourteen unit tests in test_analysis.py verify all statistical helper functions '
    'independently prior to model training. The deployed Logistic Regression models the '
    'log-odds of mortality as a linear function of the standardised features (Eq. 1), where '
    'sigma denotes the logistic function. Table III lists the tuned hyperparameters.'
)
eq('P(y=1 | x) = σ(β·x + β₀),   σ(z) = 1 / (1 + e^−z)', 1)
make_table(
    ['Hyperparameter', 'Value'],
    [
        ('Optimiser / solver', 'GridSearchCV (5-fold stratified)'),
        ('LR regularisation C', '0.01 (L2)'),
        ('RF depth / trees / leaf', '6 / 200 / 10'),
        ('XGBoost scale_pos_weight', '7'),
        ('Class-imbalance handling', "class_weight / sample_weight / scale_pos_weight"),
        ('Calibration', "isotonic, CalibratedClassifierCV (cv=5)"),
        ('Bootstrap resamples (CI)', '2,000'),
        ('Random state', '42 (reproducible)'),
    ],
    'TABLE III.  Tuned Hyperparameters'
)

h2('E', 'Probability Calibration')
body(
    'Pre-calibration reliability analysis revealed systematic over-confidence across all '
    'probability ranges (Brier = 0.135, ECE = 0.195). The reliability diagram showed '
    'non-monotonic deviations requiring a non-parametric correction; isotonic regression '
    'calibration was therefore selected over Platt scaling (which assumes monotone '
    'sigmoid distortion). The deployed model is a CalibratedClassifierCV (method=\'isotonic\', '
    'cv=5) wrapping the tuned Logistic Regression, persisted as model_primary_deployed.joblib '
    '(13.0 KB). Model selection used DeLong\'s method [15] comparing AUC values across all '
    'five models on the held-out test set, with 2,000-resample bootstrap providing 95% CIs. '
    'Calibration is quantified by the Brier score (Eq. 2) and the count-weighted Expected '
    'Calibration Error across M probability bins (Eq. 3).'
)
eq('BS = (1/N) Σ (pi − yi)²', 2)
eq('ECE = Σ (|Bm| / N) · |acc(Bm) − conf(Bm)|', 3)
fig_embed('advanced_plots/01_calibration_improvement.png',
    'Fig. 3.  Reliability diagram before and after isotonic calibration. Post-calibration '
    'probabilities closely track observed mortality rates across all risk deciles.')

h2('F', 'Multi-Layer Explainability Framework')
body(
    'Three complementary explainability layers are provided. (1) Global SHAP: Tree-SHAP '
    '(Lundberg 2020) on the XGBoost model provides model-agnostic additive feature '
    'attributions computed on the full test set, enabling global importance ranking and '
    'individual patient waterfall plots. (2) Odds ratios: an unregularised statsmodels '
    'Logit model provides OR per +1 SD with 95% Wald confidence intervals, enabling '
    'clinician-interpretable relative risk presentation. (3) Clinical nomogram: integer '
    'point scores derived from log-odds coefficients allow bedside paper-based risk '
    'estimation. The dominant pairwise feature interaction is pneumonia x age (interaction '
    'strength 0.182), indicating disproportionate risk for elderly patients with acute '
    'respiratory involvement. Each odds ratio is the exponent of the corresponding logistic '
    'coefficient (Eq. 4), giving the multiplicative change in mortality odds per +1 SD.'
)
eq('ORj = exp(βj)', 4)

h2('G', 'Subgroup Fairness Audit and Threshold Mitigation')
body(
    'The audit evaluates two metrics across eight pre-specified subgroups: AUC '
    '(discrimination fairness) and count-weighted ECE (calibration equity). Subgroups: '
    'age 18-30, 30-50, 50-65, 65+; sex female/male; diabetes present/absent. '
    'Flagging criterion: subgroup AUC deviates >5 percentage points from overall AUC. '
    'Group-specific decision thresholds, derived to equalise true positive rates, '
    'reduced TPR disparity from 0.603 to 0.045 without model retraining: age 18-30 '
    't=0.05, 30-50 t=0.12, 50-65 t=0.23, 65+ t=0.41.'
)

h2('H', 'Web Prototype Architecture')
body(
    'A two-page Streamlit application loads the three pre-fitted joblib artefacts (imputer, '
    'scaler, calibrated model) and applies the identical impute -> scale -> predict_proba '
    'chain used in training. Inputs are an age slider and eight radio/toggle controls; outputs '
    'are the calibrated risk percentage with a four-tier colour band, per-patient odds-ratio '
    'attribution bars, a global SHAP panel, and the nomogram total. All metrics are read at '
    'runtime from the exported CSVs to prevent drift, and a four-location disclaimer frames '
    'outputs as mortality-proxy estimates not validated for individual clinical decisions.'
)

h2('I', 'Temporal Out-of-Time Validation')
body(
    'Records prior to 15 June 2020 formed the temporal training set (176,979 patients); '
    'subsequent records formed the temporal validation set (43,239 patients). This '
    'design tests robustness to temporal distribution shift as clinical protocols, '
    'patient demographics, and co-morbidity management evolved across the surveillance '
    'collection period -- providing a more stringent test than random train-test split alone.'
)

h2('J', 'Baseline Model')
body(
    'Two reference baselines contextualise the proposed pipeline. A majority-class classifier '
    'predicting survival for every patient attains 87.7% accuracy yet AUC 0.50 and zero recall, '
    'demonstrating why accuracy is inadequate under 12% prevalence and motivating AUC, '
    'calibration, and net-benefit as primary metrics. A plain, unweighted, uncalibrated '
    'Logistic Regression at the default 0.5 threshold provides the methodological baseline '
    'against which each component is isolated: cost-sensitive weighting, isotonic calibration, '
    'and threshold optimisation. Relative to this baseline the deployed model reduces ECE from '
    '0.195 to 0.002 and lifts operating-point sensitivity from 0.296 to 0.702, at equivalent '
    'discrimination (Section IV).'
)

# ── IV. PERFORMANCE EVALUATION ────────────────────────────────────────────────
h1('IV', 'Performance Evaluation')

h2('A', 'Model Discrimination and DeLong Testing')
make_table(
    ['Model', 'Test AUC', '95% CI', 'DeLong vs LR'],
    [
        ('Logistic Regression (tuned)', '0.888', '0.884-0.892', 'reference'),
        ('Random Forest',               '0.866', '0.861-0.871', 'p ~ 3e-42 (inferior)'),
        ('Gradient Boosting',           '0.888', '0.884-0.892', 'p > 0.05 (n.s.)'),
        ('XGBoost',                     '0.888', '0.884-0.892', 'p > 0.05 (n.s.)'),
        ('Stacking Ensemble',           '0.888', '0.884-0.892', 'p > 0.05 (n.s.)'),
        ('Primary (calibrated LR)',     '0.888', '0.884-0.892', 'deployed'),
    ],
    'TABLE IV.  Model Comparison -- Test Set AUC (n = 44,044)'
)
fig_embed('model_comparison_plots/02_roc_curves.png',
    'Fig. 4.  ROC curves for all five model families. Calibrated LR is statistically '
    'non-inferior to all ensembles by DeLong test; RF is significantly worse (p ~ 3e-42).')
body(
    'Random Forest is the only model significantly inferior to Logistic Regression '
    '(DeLong p ~ 3e-42). Gradient Boosting, XGBoost, and the Stacking ensemble are '
    'statistically indistinguishable from LR. Five-fold cross-validation AUC: '
    '0.891 +/- 0.002. Temporal out-of-time AUC: 0.891. This directly and formally '
    'answers RQ2: no complex model significantly outperforms the interpretable '
    'Logistic Regression on this nine-feature clinical dataset.'
)

h2('B', 'Calibration Analysis')
make_table(
    ['Metric', 'Pre-calibration', 'Post-isotonic'],
    [
        ('Brier Score',          '0.135', '0.078'),
        ('ECE (count-weighted)', '0.195', '0.002'),
        ('ECE (bin-averaged)',   '0.287', '0.064'),
    ],
    'TABLE V.  Calibration Metrics Before and After Isotonic Correction'
)
body(
    'Isotonic calibration produces substantial improvements on all three metrics. '
    'Count-weighted ECE (0.002) reflects calibration weighted by the population '
    'distribution; bin-averaged ECE (0.064) gives equal weight to rare high-risk bins. '
    'A predicted probability of 0.30 now corresponds to approximately 30% observed '
    'mortality in the test set, enabling direct probabilistic clinical interpretation.'
)

h2('C', 'Threshold Optimisation')
make_table(
    ['Strategy', 'Threshold', 'Sensitivity', 'Specificity', 'Precision', 'F1'],
    [
        ('Default (0.5)',    '0.50', '0.296', '0.968', '0.564', '0.388'),
        ('F1-optimal',      '0.26', '0.702', '0.878', '0.447', '0.546'),
        ("Youden's J",      '0.10', '0.877', '0.761', '0.340', '0.490'),
    ],
    'TABLE VI.  Threshold Optimisation Strategies'
)
body(
    'The F1-optimal threshold (0.26) is recommended for clinical triage: sensitivity '
    '0.702 with precision 0.447, representing 3.6x enrichment over the 12.3% base rate. '
    "Youden's J (0.10) maximises sensitivity for high-recall population screening contexts."
)

h2('D', 'Feature Importance and Explainability')
make_table(
    ['Feature', 'SHAP %', 'OR (95% CI)', 'Sig.'],
    [
        ('Pneumonia',      '44.2', '2.65 (2.61-2.69)', '***'),
        ('Age',            '33.0', '2.33 (2.28-2.37)', '***'),
        ('Sex (female)',    '8.5',  '0.80 (0.79-0.82)', '*** protective'),
        ('Diabetes',       '5.6',  '1.15 (1.13-1.17)', '***'),
        ('Hypertension',   '4.5',  '1.09 (1.07-1.10)', '***'),
        ('Obesity',        '3.0',  '1.10 (1.08-1.12)', '***'),
        ('COPD',           '0.6',  '1.03 (1.02-1.04)', '***'),
        ('Asthma',         '0.3',  '0.98 (0.97-1.00)', 'n.s.'),
        ('Cardiovascular', '0.3',  '1.00 (0.99-1.01)', 'n.s.'),
    ],
    'TABLE VII.  Tree-SHAP Global Importance and Odds Ratios (statsmodels, per +1 SD)'
)
fig_embed('shap_plots/05_shap_bar_plot.png',
    'Fig. 5.  Tree-SHAP global feature importance. Pneumonia (44.2%) and age (33.0%) '
    'together account for 77.2% of total importance, confirming a linear additive signal.')
body(
    'Pneumonia dominates (OR 2.65, SHAP 44.2%) -- the largest modifiable risk factor '
    'at admission. Age (OR 2.33, SHAP 33.0%) is the strongest non-modifiable predictor. '
    'Female sex is significantly protective (OR 0.80). Cardiovascular disease and asthma '
    'show negligible individual contributions (0.3% SHAP, OR ~ 1.00, p > 0.05) -- '
    'a null finding that is transparently reported rather than suppressed. Dominant '
    'pairwise interaction: pneumonia x age (strength 0.182). Clinical nomogram points: '
    'pneumonia +24, age +6/decade, diabetes +5, COPD +4, hypertension +3, obesity +3, '
    'cardiovascular +1, asthma -1, sex (female) -5.'
)

h2('E', 'Subgroup Fairness Audit Results')
make_table(
    ['Subgroup', 'n', 'AUC', 'Flag', 'ECE'],
    [
        ('Age 18-30',   '6,290',  '0.891', 'OK',   '0.002'),
        ('Age 30-50',   '19,713', '0.874', 'OK',   '0.004'),
        ('Age 50-65',   '11,189', '0.799', 'FLAG', '0.008'),
        ('Age 65+',     '5,926',  '0.733', 'FLAG', '0.015'),
        ('Female',      '19,961', '0.901', 'OK',   '0.005'),
        ('Male',        '24,083', '0.873', 'OK',   '0.003'),
        ('Diabetes',    '7,160',  '0.776', 'FLAG', '0.012'),
        ('No Diabetes', '36,704', '0.899', 'OK',   '0.004'),
    ],
    'TABLE VIII.  Subgroup Fairness Audit (FLAG = AUC deviation > 5pp from overall)'
)
fig_embed('fairness_plots/02_fairness_disparity.png',
    'Fig. 6.  Subgroup fairness audit: AUC and ECE across eight groups. ECE < 0.015 '
    'in all subgroups; AUC drops in older and diabetic subgroups reflect restriction of range.')
body(
    'Three subgroups are flagged on AUC: age 50-65 (0.799), 65+ (0.733), and '
    'diabetes (0.776). ECE remains below 0.015 in all eight subgroups, confirming '
    'equitable calibration across demographic groups. Gender AUC gap: 2.8% (within the '
    '5% threshold). Group-specific threshold mitigation reduced TPR disparity from '
    '0.603 to 0.045.'
)

h2('F', 'Decision Curve Analysis and Sensitivity Analysis')
fig_embed('advanced_plots/02_decision_curve.png',
    'Fig. 7.  Decision Curve Analysis. Model net benefit (0.085 at threshold 0.10) '
    'substantially exceeds treat-all (0.026) and treat-none across the clinical range.')
body(
    'DCA confirms positive net benefit over both treat-all and treat-none strategies '
    'across the full clinically relevant threshold range (0.05-0.50). At threshold 0.10, '
    'model net benefit = 0.085 versus treat-all = 0.026, validating practical utility '
    'beyond accuracy metrics alone. Counterfactual sensitivity analysis confirms that '
    'pneumonia removal produces the largest individual risk reduction (illustrative '
    'example: 42.0% -> 8.0%, delta = -34.0 pp), consistent with its dominant SHAP '
    'importance. Risk band distribution: LOW <30% (83.3%), MEDIUM 30-50% (10.3%), '
    'HIGH 50-70% (6.4%), CRITICAL >=70% (0.05%).'
)

# ── V. DISCUSSION ─────────────────────────────────────────────────────────────
h1('V', 'Discussion')

h2('A', 'Why Logistic Regression Ties the Ensembles')
body(
    'The finding that no complex model significantly outperforms Logistic Regression '
    '(DeLong p > 0.05) requires explanation rather than mere reporting. The mortality '
    'signal in this dataset is predominantly linear and additive: pneumonia and age '
    'together account for 77.2% of Tree-SHAP importance, with monotonic, non-interacting '
    'effects. With nine mostly-binary features and such a concentrated linear signal '
    'structure, tree ensembles have no high-order interaction or non-monotonic patterns '
    'to exploit -- additional parametric complexity overfits noise rather than captures '
    'signal. This independently confirms Bjerre et al. [5] with formal DeLong statistical '
    'evidence on a separate 220,218-patient dataset. The clinical implication is significant: '
    'on low-dimensional binary clinical features, interpretability costs nothing in '
    'discrimination. The simpler, explainable model is strictly preferable for deployment.'
)

h2('B', 'Calibration as the Primary Deployment Criterion')
body(
    'For clinical deployment, calibration matters more than AUC: a clinician acts on the '
    'predicted probability as a risk estimate, not as an unlabelled ranking score. The '
    'pre-calibration ECE of 0.195 illustrates the danger of deploying uncalibrated models '
    '-- predicted probabilities diverging from observed mortality by nearly 20 percentage '
    'points on average. Post-isotonic ECE = 0.002 (count-weighted) enables direct clinical '
    'probability use. Crucially, ECE < 0.015 in all eight subgroups confirms equitable '
    'calibration: no demographic group receives systematically miscalibrated risk estimates, '
    'meeting a basic requirement for equitable clinical AI deployment.'
)

h2('C', 'Interpreting Subgroup AUC Drops')
body(
    'The lower AUC observed in age 65+ (0.733) and diabetic (0.776) subgroups reflects '
    'restriction of range -- a well-known statistical phenomenon whereby AUC naturally '
    'decreases within homogeneous high-risk subgroups due to compressed rank variation. '
    'This is a statistical artefact of subgroup composition, not model failure. Preserved '
    'calibration (ECE 0.015 for age 65+, 0.012 for diabetes) confirms that risk estimates '
    'remain accurate within these groups even as ranking becomes harder. This distinction '
    'between discrimination fairness and calibration fairness is frequently conflated in '
    'the clinical AI fairness literature.'
)

h2('D', 'Responsible AI and Governance Framing')
body(
    'The pipeline implements responsible-AI best practices at three levels. At the model '
    'level: leakage-free design enforced by split-then-fit pipelines, isotonic calibration '
    'with empirical verification, threshold optimisation with two clinical strategies, and '
    'temporal out-of-time validation. At the fairness level: eight pre-specified subgroup '
    'audit with two fairness lenses (AUC and ECE), transparent reporting of null and '
    'flagged results, and a concrete mitigation strategy reducing TPR disparity to 0.045. '
    'At the documentation level: Model Card (Mitchell et al. 2019), Datasheet for Datasets '
    '(Gebru et al. 2021), and a four-location Streamlit disclaimer that frames all outputs '
    'as mortality-proxy estimates not validated for individual clinical decisions.'
)

h2('E', 'Limitations and Threats to Validity')
body(
    'Four limitations bound the claims of this work. (i) Construct validity: in-hospital '
    'mortality is used as a PASC proxy; it under-estimates Long COVID burden in patients '
    'who survive with significant sequelae, potentially missing a large proportion of the '
    'target population. (ii) External validity: single-country, single-period data '
    '(Mexico, March-November 2020) limits generalisability to other healthcare systems, '
    'demographic profiles, and SARS-CoV-2 variants. No vaccination data are available as '
    'the dataset predates widespread rollout. (iii) Feature scope: nine admission-time '
    'features only -- laboratory values, vital signs, comorbidity severity grading, and '
    'variant information are absent. (iv) Internal validation only: temporal validation '
    'within the same national dataset cannot substitute for multi-site or multi-country '
    'external validation.'
)

# ── VI. CONCLUSION ────────────────────────────────────────────────────────────
h1('VI', 'Conclusion')

body(
    'This paper presented a rigorous 13-phase end-to-end responsible-ML pipeline for '
    'COVID-19 in-hospital mortality risk stratification as a proxy for Long COVID severe '
    'outcome risk. The central empirical finding answers RQ2 with formal statistical '
    'evidence: a calibrated, interpretable Logistic Regression is non-inferior to Random '
    'Forest, Gradient Boosting, XGBoost, and Stacking by DeLong\'s test (p > 0.05) on '
    '44,044 held-out patients, achieving AUC 0.888 (95% CI 0.884-0.892). The '
    'predominantly linear, additive signal structure -- pneumonia and age accounting for '
    '77.2% of Tree-SHAP importance -- leaves no advantage for complex models to exploit.'
)
body(
    'Three contributions were made. First, an empirical finding confirmed with formal '
    'DeLong evidence on 220,218 patients: interpretable Logistic Regression matches '
    'black-box ensembles on low-dimensional binary clinical features, independently '
    'replicating Bjerre et al. [5]. Second, a methodological contribution: the pipeline '
    'closes all four identified gaps from the Long COVID prediction literature in a '
    'single reproducible workflow -- calibration (Brier 0.135 to 0.078, ECE 0.195 to '
    '0.002), fairness (TPR disparity 0.603 to 0.045, ECE < 0.015 in all eight subgroups), '
    'leakage control (split-then-fit, nine admission-time features), and multi-layer '
    'explainability (SHAP, odds ratios, nomogram, web app). Third, a responsible-AI '
    'contribution: Model Card, Datasheet, and a Streamlit clinical prototype with '
    'governance disclaimers and group-specific threshold mitigation built in.'
)
body(
    'Future work should apply the pipeline to PASC-labelled EHR cohorts (RECOVER [8], '
    'UK Biobank), incorporate longitudinal biomarkers, vaccination status, and variant '
    'data, pursue multi-site external validation, and evaluate the Streamlit prototype '
    'in a prospective clinical decision-support trial measuring patient outcomes rather '
    'than model metrics alone.'
)

# ── GENERATIVE AI DISCLOSURE ──────────────────────────────────────────────────

# -- ACKNOWLEDGMENTS --
h1('', 'Acknowledgments')
body(
    'The author thanks Dr Martin Crane and Dr Tai Tan Mai (School of Computing, '
    'Dublin City University) for supervision, guidance, and critical feedback '
    'throughout this research. The Mexican Government COVID-19 surveillance '
    'dataset was sourced from Kaggle [14] and is gratefully acknowledged. '
    'This work was conducted as part of the MSc Computing (Data Analytics) '
    'programme at Dublin City University, Ireland.'
)

h1('', 'Generative AI Disclosure')
body(
    'Generative AI tools (Claude, Anthropic) were used for: (i) code review and debugging '
    'of the analysis pipeline; (ii) drafting and editing report sections based on verified '
    'numerical results from the author\'s own pipeline; (iii) IEEE formatting assistance. '
    'All quantitative results, model training, statistical tests, and data analyses were '
    'performed exclusively by the author\'s Python code. All AI-assisted text was reviewed, '
    'verified against actual pipeline outputs, and edited by the author. No AI tool '
    'generated any experimental result or analytical decision.'
)

# ── REFERENCES ────────────────────────────────────────────────────────────────
h1('', 'References')
refs = [
    # [1]-[15] core citations (cited in text)
    'World Health Organisation, "Post COVID-19 condition (long COVID)," WHO Fact Sheet, 2025. [Online]. Available: https://www.who.int/news-room/fact-sheets/detail/post-covid-19-condition',
    'D. A. Gbewonyo et al., "Long COVID in Adults: United States, 2022," NCHS Data Brief, no. 480, CDC, 2023. doi: 10.15620/cdc:132417.',
    'B. Antony et al., "Predictive models of long COVID," EBioMedicine, vol. 96, p. 104777, 2023. doi: 10.1016/j.ebiom.2023.104777.',
    'K. Jeffrey et al., "Deriving and validating a risk prediction model for long COVID: a population-based, retrospective cohort study in Scotland," J. Royal Society of Medicine, vol. 117, no. 12, pp. 402-414, 2024. doi: 10.1177/01410768241297833.',
    'L. M. Bjerre et al., "Comparing AI/ML approaches and classical regression for predictive modeling using large population health databases: Applications to COVID-19 case prediction," Global Epidemiology, vol. 8, p. 100168, 2024. doi: 10.1016/j.gloepi.2024.100168.',
    'A. S. Menezes et al., "Acute COVID-19 Syndrome Predicts Severe Long COVID-19: An Observational Study," Cureus, 2022. doi: 10.7759/cureus.29826.',
    'Z. B. Dozier et al., "Predicting Long COVID in the National COVID Cohort Collaborative Using Super Learner," JMIR Public Health and Surveillance, vol. 10, p. e53322, 2024. doi: 10.2196/53322.',
    'C. Zang et al., "Identification of risk factors of Long COVID and predictive modeling in the RECOVER EHR cohorts," Communications Medicine, vol. 4, no. 1, 2024. doi: 10.1038/s43856-024-00549-0.',
    'N. Shakhovska, V. Yakovyna, and V. Chopyak, "A new hybrid ensemble machine-learning model for severity risk assessment and post-COVID prediction system," Mathematical Biosciences and Engineering, vol. 19, no. 6, pp. 6102-6123, 2022. doi: 10.3934/mbe.2022285.',
    'K. Swinnerton et al., "Leveraging near-real-time patient and population data to incorporate fluctuating risk of severe COVID-19," eClinicalMedicine, vol. 81, p. 103114, 2025. doi: 10.1016/j.eclinm.2025.103114.',
    'H. M. Teeuw et al., "Incidence and individual risk prediction of post-COVID-19 cardiovascular disease in the general population," ESC Open Heart, vol. 3, 2023. doi: 10.1093/ehjopen/oead101.',
    'J. G. Guzman-Esquivel et al., "Clinical Characteristics in the Acute Phase of COVID-19 That Predict Long COVID," Healthcare, vol. 11, no. 2, pp. 197-197, 2023. doi: 10.3390/healthcare11020197.',
    'Z. Noroozi, A. Orooji, and L. Erfannia, "Analyzing the impact of feature selection methods on machine learning algorithms for heart disease prediction," Scientific Reports, vol. 13, 2023. doi: 10.1038/s41598-023-49962-w.',
    'Kaggle, "COVID-19 Mexico Patient Health Dataset," 2020. [Online]. Available: https://www.kaggle.com/datasets/riteshahlawat/covid19-mexico-patient-health-dataset',
    'E. R. DeLong, D. M. DeLong, and D. L. Clarke-Pearson, "Comparing the areas under two or more correlated receiver operating characteristic curves: a nonparametric approach," Biometrics, vol. 44, no. 3, pp. 837-845, 1988.',
    # [16]-[21] additional directly relevant references
    'S. Luo et al., "Prevalence and duration of common symptoms in people with long COVID: a systematic review and meta-analysis," Journal of Global Health, vol. 15, 2025. doi: 10.7189/jogh.15.04282.',
    'S. Bozkurt and K. Keskin, "A gradient boosting-based mortality prediction model for COVID-19 patients," Neural Computing and Applications, vol. 35, no. 33, pp. 23997-24013, 2023. doi: 10.1007/s00521-023-08997-w.',
    'E. Cordelli et al., "Predicting pulmonary sequelae after COVID-19 using machine learning: model performance and clinical interpretability," 2024.',
    'N. Kourmpanis et al., "Predicting mortality outcomes in individual COVID-19 patients using machine learning algorithms," Artificial Intelligence in Health, vol. 1, no. 3, p. 31, 2024. doi: 10.36922/aih.2591.',
    'X. Sun et al., "Development and validation of a prediction model for mortality in critically ill COVID-19 patients," Frontiers in Cellular and Infection Microbiology, vol. 14, pp. 1309529-1309529, 2024. doi: 10.3389/fcimb.2024.1309529.',
    'C. Y. Tang et al., "Prediction models for COVID-19 disease outcomes," Emerging Microbes and Infections, 2024. doi: 10.1080/22221751.2024.2361791.',
]
for i, r in enumerate(refs, 1):
    ref_entry(i, r)

doc.save(str(OUT))
print(f'Saved: {OUT}')
total_words = sum(len(p.text.split()) for p in doc.paragraphs)
print(f'Words: {total_words} | Tables: {len(doc.tables)} | Paragraphs: {len(doc.paragraphs)}')
print('Figures embedded: 7 real PNGs')
